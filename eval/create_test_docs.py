"""
Generate test fixture documents for RAG evaluation.

Creates:
  docs/test-documents/company-policies.docx  — contradictory timelines + multi-hop
  docs/test-documents/tech-reference.pdf     — overlapping contexts (Project Apollo)

Usage:
  pip install python-docx fpdf2
  python eval/create_test_docs.py
"""

import pathlib

OUT_DIR = pathlib.Path(__file__).parent.parent / "docs" / "test-documents"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def create_docx():
    from docx import Document

    doc = Document()
    doc.add_heading("Company Policy Document", 0)

    doc.add_heading("REMOTE WORK POLICY 2024", level=1)
    doc.add_paragraph(
        "All employees are required to work from the office 4 days per week. "
        "Remote work is permitted on Fridays only. "
        "Employees must notify their manager at least 24 hours in advance when working remotely."
    )

    doc.add_heading("REMOTE WORK POLICY 2026", level=1)
    doc.add_paragraph(
        "The company adopts a remote-first approach. "
        "Employees are authorized to work from home every day of the week. "
        "Office attendance is encouraged but never mandatory. "
        "All team meetings will be conducted via video conferencing by default."
    )

    doc.add_heading("OFFICE DIRECTORY", level=1)
    doc.add_paragraph(
        "John Smith is a Senior Software Engineer. John Smith works in Office A. "
        "Office A is located in Lviv, Ukraine. "
        "Office A houses the Engineering and Product teams."
    )

    doc.add_heading("EQUIPMENT POLICY", level=1)
    doc.add_paragraph(
        "All employees receive a company laptop upon joining. "
        "Laptops must be returned within 5 business days of employment termination. "
        "Employees may request ergonomic peripherals through the IT portal."
    )

    out = OUT_DIR / "company-policies.docx"
    doc.save(out)
    print(f"Created: {out}")


def create_pdf():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Tech Project Reference", ln=True, align="C")
    pdf.ln(5)

    sections = [
        (
            "Project Apollo Phase 1",
            "Project Apollo Phase 1 uses AWS (Amazon Web Services) as its cloud infrastructure. "
            "The phase focuses on initial data migration and API development. "
            "Key services include Amazon S3 for storage, EC2 for compute, and RDS for the relational database. "
            "Phase 1 was completed in Q2 2024 with 99.9% uptime achieved.",
        ),
        (
            "Project Apollo Phase 2",
            "Project Apollo Phase 2 migrates the infrastructure to GCP (Google Cloud Platform). "
            "This phase introduces machine learning pipelines and BigQuery integration. "
            "Key services include Cloud Run for containerized workloads, Vertex AI for model training, "
            "and Pub/Sub for event streaming. Phase 2 is scheduled for completion in Q4 2025.",
        ),
        (
            "Project Hermes",
            "Project Hermes is an internal developer tooling initiative hosted on Azure. "
            "It provides CI/CD pipelines using Azure DevOps and artifact management via Azure Container Registry. "
            "The project is maintained by the Platform Engineering team.",
        ),
    ]

    for title, body in sections:
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, title, ln=True)
        pdf.set_font("Helvetica", size=11)
        pdf.multi_cell(0, 6, body)
        pdf.ln(4)

    out = OUT_DIR / "tech-reference.pdf"
    pdf.output(str(out))
    print(f"Created: {out}")


if __name__ == "__main__":
    create_docx()
    create_pdf()
    print("Done.")
